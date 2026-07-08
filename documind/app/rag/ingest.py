"""Multi-format ingestion with hierarchical (parent->child) chunking + numeric check.
Parses pdf/docx/xlsx/csv/txt/images -> parent chunks -> child (leaf) chunks with
parent_text context -> embed leaves -> store. Verifies numeric preservation."""
import re
import uuid
from pathlib import Path
import fitz  # pymupdf
import docx
import openpyxl
from app import config
from app.rag import embed, store, llm

IMAGE_EXT = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}
MIME = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".webp": "image/webp", ".gif": "image/gif", ".bmp": "image/bmp"}
_NUM = re.compile(r"[-+]?\$?\d{1,3}(?:,\d{3})+(?:\.\d+)?%?|[-+]?\$?\d+(?:\.\d+)?%?")


def _sections(path: Path) -> list[dict]:
    ext = path.suffix.lower()
    out: list[dict] = []
    if ext == ".pdf":
        doc = fitz.open(path)
        for i, page in enumerate(doc):
            t = page.get_text("text").strip()
            if t:
                out.append({"text": t, "location": f"p.{i + 1}"})
        doc.close()
    elif ext == ".docx":
        d = docx.Document(str(path))
        paras = [p.text for p in d.paragraphs if p.text.strip()]
        for tbl in d.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paras.append(" | ".join(cells))
        out.append({"text": "\n".join(paras), "location": "document"})
    elif ext in (".xlsx", ".xlsm"):
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        for ws in wb.worksheets:
            rows, header = [], None
            for r in ws.iter_rows(values_only=True):
                vals = [("" if v is None else str(v)) for v in r]
                if not any(vals):
                    continue
                if header is None:
                    header = vals
                    continue
                pairs = [f"{h}: {v}" for h, v in zip(header, vals) if v]
                rows.append("; ".join(pairs) if pairs else " | ".join(vals))
            if header:
                rows.insert(0, "Columns: " + ", ".join(h for h in header if h))
            if rows:
                out.append({"text": "\n".join(rows), "location": f"sheet:{ws.title}"})
        wb.close()
    elif ext in (".csv", ".txt", ".md"):
        out.append({"text": path.read_text(errors="ignore"), "location": "document"})
    elif ext in IMAGE_EXT:
        text, _, _ = llm.image_to_text(path.read_bytes(), MIME.get(ext, "image/png"))
        out.append({"text": text, "location": "image"})
    else:
        out.append({"text": path.read_text(errors="ignore"), "location": "document"})
    return [s for s in out if s["text"].strip()]


def _split(text: str, size: int, overlap: int) -> list[str]:
    text = text.strip()
    if len(text) <= size:
        return [text] if text else []
    out, start = [], 0
    while start < len(text):
        end = start + size
        if end < len(text):
            brk = text.rfind("\n", start + size // 2, end)
            if brk == -1:
                brk = text.rfind(" ", start + size // 2, end)
            if brk != -1:
                end = brk
        out.append(text[start:end].strip())
        start = max(end - overlap, end)
    return [c for c in out if c]


def _hierarchical(sections: list[dict]) -> list[dict]:
    """parent chunks -> child (leaf) chunks carrying their parent's text as context."""
    leaves = []
    idx = 0
    for sec in sections:
        for parent in _split(sec["text"], config.PARENT_CHARS, 0):
            children = _split(parent, config.CHILD_CHARS, config.CHILD_OVERLAP) or [parent]
            for child in children:
                leaves.append({"text": child, "parent_text": parent,
                               "location": sec["location"], "chunk_index": idx})
                idx += 1
    return leaves


def _numeric_check(source: str, chunk_texts: list[str]) -> dict:
    src = set(_NUM.findall(source))
    if not src:
        return {"total": 0, "preserved": 0, "missing": 0, "score": 1.0}
    joined = " ".join(chunk_texts)
    kept = {n for n in src if n in joined}
    return {"total": len(src), "preserved": len(kept), "missing": len(src - kept),
            "score": round(len(kept) / len(src), 3)}


def ingest_file(user_id: int, path: Path, filename: str, folder: str) -> dict:
    sections = _sections(path)
    leaves = _hierarchical(sections)
    if not leaves:
        return {"doc_id": None, "chunks": 0, "embed_tokens": 0, "numeric": None}
    vectors, tokens = embed.embed([c["text"] for c in leaves], input_type="document")
    doc_id = str(uuid.uuid4())
    store.add_chunks(user_id, doc_id, filename, folder, leaves, vectors)
    numeric = _numeric_check(" ".join(s["text"] for s in sections), [c["text"] for c in leaves])
    return {"doc_id": doc_id, "chunks": len(leaves), "embed_tokens": tokens, "numeric": numeric}
